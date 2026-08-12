"""Fill the official final-project submission form from league evidence.

Game rows and totals are DERIVED (never typed by hand) from the league ledger
``results/counted_series.json`` plus the per-game ``result_*.json`` artifacts
(see :mod:`tools.submission_builder.rows`); identity/self-score fields come from
a JSON data file kept OUTSIDE the repos.
Filling goes through :mod:`tools.pdf_parser.docx_form`, which only appends into
existing runs — the template's fields never move (a hard course requirement).
"""

from __future__ import annotations

import json
from pathlib import Path

from tools.pdf_parser.docx_form import docx_to_pdf, fill_docx_form
from tools.submission_builder.rows import (
    IL as IL,
)
from tools.submission_builder.rows import (
    compute_totals,
    load_games,
    unfilled_placeholders,
)

#: Paragraph indices of the official uoh-rl07 final-project template.
FIELD_PARAGRAPHS = {
    "group_id": 1,
    "self_score": 2,
    "cop_repo": 3,
    "thief_repo": 4,
    "agent_email": 5,
    "student1_id": 7,
    "student2_id": 11,
    "legal_games": 14,
    "points": 15,
    "won": 16,
    "lost": 17,
    "drawn": 18,
    "bonus": 19,
}
NAME_PARAGRAPHS = {
    8: ("student1", "en"),
    9: ("student1", "he"),
    12: ("student2", "en"),
    13: ("student2", "he"),
}


def build_submission(
    template: str | Path,
    data_file: str | Path,
    results_dir: str | Path,
    out_dir: str | Path,
    to_pdf: bool = True,
) -> dict:
    """Fill the form and (optionally) convert to PDF. Returns paths + warnings."""
    data = json.loads(Path(data_file).read_text(encoding="utf-8"))
    games = load_games(results_dir, data["group_id"])
    totals = compute_totals(games)
    s1, s2 = data["student1"], data["student2"]
    values = {
        FIELD_PARAGRAPHS["group_id"]: data["group_id"],
        FIELD_PARAGRAPHS["self_score"]: data["self_score"],
        FIELD_PARAGRAPHS["cop_repo"]: data["cop_repo"],
        FIELD_PARAGRAPHS["thief_repo"]: data["thief_repo"],
        FIELD_PARAGRAPHS["agent_email"]: data["agent_email"],
        FIELD_PARAGRAPHS["student1_id"]: s1["id_card"],
        FIELD_PARAGRAPHS["student2_id"]: s2["id_card"],
        FIELD_PARAGRAPHS["bonus"]: data["bonus_eligibility"],
        **{
            FIELD_PARAGRAPHS[k]: str(totals[k])
            for k in ("legal_games", "points", "won", "lost", "drawn")
        },
    }
    emails = data["opponent_agent_emails"]
    rows = [
        [
            str(i),
            g["date"],
            g["start"],
            g["end"],
            g["opponent"],
            str(g["us"]),
            str(g["them"]),
            str(g["declared"]),
            emails.get(g["opponent"], ""),
        ]
        for i, g in enumerate(games, start=1)
    ]
    yy = str(data["exercise_number"])
    stem = f"{data['group_id']}-ex{'XX' if '<FILL' in yy else yy}"
    out_dir = Path(out_dir)
    out_docx = fill_docx_form(template, out_dir / f"{stem}.docx", values, table_rows=rows)
    _fill_names(out_docx, {"student1": s1, "student2": s2})
    result = {"docx": out_docx, "pdf": None, "unfilled": unfilled_placeholders(data)}
    if to_pdf:
        result["pdf"] = docx_to_pdf(out_docx)
    return result


def _fill_names(docx_path: Path, students: dict) -> None:
    from docx import Document

    doc = Document(str(docx_path))
    for index, (who, lang) in NAME_PARAGRAPHS.items():
        student = students[who]
        first = student[f"first_{lang}"]
        last = student[f"last_{lang}"]
        text = doc.paragraphs[index].text
        head, _, tail = text.partition("Last name")
        for run in list(doc.paragraphs[index].runs):
            run.text = ""
        doc.paragraphs[index].add_run(f"{head.rstrip()} {first}    Last name{tail} {last}")
    doc.save(str(docx_path))
