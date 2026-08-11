"""Fill Word-template forms and convert them to PDF.

Course submissions must be built FROM the official .docx template with every
field kept in place, so this module never creates layout — it only appends
values into existing paragraphs/tables and then converts the result to PDF.

Requires ``python-docx`` (and ``docx2pdf`` + MS Word for the PDF step); both are
optional extras — importing this module without them raises a clear error only
when the functions are called.
"""

from __future__ import annotations

from pathlib import Path

from tools.pdf_parser.reader import PdfParserError


def fill_docx_form(
    template: str | Path,
    output: str | Path,
    paragraph_values: dict[int, str],
    table_rows: list[list[str]] | None = None,
    table_index: int = 0,
    first_data_row: int = 1,
) -> Path:
    """Append ``paragraph_values[i]`` to paragraph ``i`` and fill table rows.

    Values are appended into each paragraph's last run so the template's field
    text, ordering, and positions are never altered. Table rows are written
    cell-by-cell starting at ``first_data_row`` (row 0 is usually the header).
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise PdfParserError("python-docx is required: uv run --with python-docx ...") from exc
    template = Path(template)
    if not template.is_file():
        raise PdfParserError(f"template not found: {template}")
    doc = Document(str(template))
    paragraphs = doc.paragraphs
    for index, value in paragraph_values.items():
        if not 0 <= index < len(paragraphs):
            raise PdfParserError(f"paragraph index {index} out of range 0..{len(paragraphs) - 1}")
        paragraph = paragraphs[index]
        if paragraph.runs:
            paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip() + " " + str(value)
        else:
            paragraph.add_run(" " + str(value))
    if table_rows:
        if table_index >= len(doc.tables):
            raise PdfParserError(f"no table {table_index} in template")
        table = doc.tables[table_index]
        for offset, values in enumerate(table_rows):
            row = table.rows[first_data_row + offset]
            if len(values) != len(row.cells):
                raise PdfParserError(
                    f"row {offset}: {len(values)} values for {len(row.cells)} cells"
                )
            for cell, value in zip(row.cells, values, strict=True):
                cell.text = str(value)
    output = Path(output)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))
    return output


def replace_docx_paragraph(document, index: int, text: str) -> None:
    """Rewrite one paragraph's text in place (all runs collapsed into one)."""
    paragraph = document.paragraphs[index]
    for run in list(paragraph.runs):
        run.text = ""
    paragraph.add_run(text)


def docx_to_pdf(docx_path: str | Path, pdf_path: str | Path | None = None) -> Path:
    """Convert a .docx to PDF via MS Word (docx2pdf). Returns the PDF path."""
    docx_path = Path(docx_path)
    target = Path(pdf_path) if pdf_path else docx_path.with_suffix(".pdf")
    try:
        from docx2pdf import convert
    except ImportError as exc:  # pragma: no cover - environment-dependent
        raise PdfParserError("docx2pdf is required: uv run --with docx2pdf ...") from exc
    convert(str(docx_path), str(target))
    if not target.is_file():
        raise PdfParserError(f"conversion produced no file: {target}")
    return target
