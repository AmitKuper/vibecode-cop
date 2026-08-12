"""tools/pdf_parser docx_form — fill a .docx form without moving its fields."""

from __future__ import annotations

import pytest

from tools.pdf_parser import PdfParserError


def _docx_available():
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed in this venv")
def test_fill_docx_form_appends_without_moving_fields(tmp_path):
    from docx import Document

    from tools.pdf_parser.docx_form import fill_docx_form

    template = tmp_path / "form.docx"
    doc = Document()
    doc.add_paragraph("Group ID code:")
    doc.add_paragraph("Self score:")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "opponent"
    table.rows[0].cells[1].text = "score"
    doc.save(str(template))

    out = tmp_path / "filled.docx"
    fill_docx_form(
        template,
        out,
        {0: "vibecode", 1: "90"},
        table_rows=[["imreeyal", "90:30"], ["anrbj666", "35:75"]],
    )
    filled = Document(str(out))
    assert filled.paragraphs[0].text == "Group ID code: vibecode"
    assert filled.paragraphs[1].text == "Self score: 90"
    assert filled.tables[0].rows[1].cells[0].text == "imreeyal"
    assert filled.tables[0].rows[2].cells[1].text == "35:75"


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed in this venv")
def test_fill_docx_form_rejects_bad_indices(tmp_path):
    from docx import Document

    from tools.pdf_parser.docx_form import fill_docx_form

    template = tmp_path / "form.docx"
    doc = Document()
    doc.add_paragraph("only one line")
    doc.save(str(template))
    with pytest.raises(PdfParserError, match="out of range"):
        fill_docx_form(template, tmp_path / "x.docx", {5: "nope"})
    with pytest.raises(PdfParserError, match="template not found"):
        fill_docx_form(tmp_path / "missing.docx", tmp_path / "x.docx", {0: "v"})
