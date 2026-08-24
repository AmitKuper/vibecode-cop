"""tools/submission_builder — ledger-derived rows/totals and the form fill."""

from __future__ import annotations

import json

import pytest

from tools.pdf_parser.reader import PdfParserError
from tools.submission_builder import compute_totals, load_games, unfilled_placeholders


def _entry(op, winner, us, them):
    gid, ts = f"{op}-vs-vibecode", {"vibecode": us, op: them}
    return {"game_id": gid, "opponent": op, "winner_group": winner, "total_score": ts}


def _write_fixture_results(tmp_path, with_draw=False):
    series = [_entry("oppa", "oppa", 35, 75), _entry("oppb", "vibecode", 90, 30)]
    days = [("oppa", "08"), ("oppb", "10")]
    if with_draw:
        series.append(_entry("oppc", None, 47, 47))
        days.append(("oppc", "20"))
    (tmp_path / "counted_series.json").write_text(
        json.dumps({"group_id": "vibecode", "counted_games_played": len(series), "series": series}),
        encoding="utf-8",
    )
    for name, day in days:
        final = {"games_played_including_this": {name: 2, "vibecode": 1}}
        if name == "oppc":
            final["series_tie"] = True
        p = f"2026-08-{day}T"
        subs = [
            {"started_at": p + "19:14:06+00:00", "ended_at": p + "19:15:00+00:00"},
            {"started_at": p + "19:15:30+00:00", "ended_at": p + "19:16:13+00:00"},
        ]
        (tmp_path / f"result_{name}-vs-vibecode.json").write_text(
            json.dumps({"final_result": final, "sub_games": subs}), encoding="utf-8"
        )


def test_load_games_derives_rows_from_ledger(tmp_path):
    _write_fixture_results(tmp_path)
    games = load_games(tmp_path)
    assert [g["opponent"] for g in games] == ["oppa", "oppb"]
    assert games[0]["us"] == 35 and games[0]["them"] == 75 and not games[0]["won"]
    assert games[1]["us"] == 90 and games[1]["won"]
    # 19:14 UTC == 22:14 league time (Asia/Jerusalem)
    assert games[0]["start"] == "22:14" and games[0]["end"] == "22:16"
    # handshake-declared count = games_played_including_this - 1 (BEFORE us)
    assert games[0]["declared"] == 1


def test_compute_totals(tmp_path):
    _write_fixture_results(tmp_path)
    totals = compute_totals(load_games(tmp_path))
    assert totals == {"legal_games": 2, "points": 125, "won": 1, "lost": 1, "drawn": 0}


def test_a_drawn_series_counts_as_drawn_not_lost(tmp_path):
    """series_tie lives in the RESULT's final_result, not the ledger row —
    reading it from the entry silently filed every draw as a loss (caught on
    the real 10-game form: lost 3 / drawn 0 instead of lost 1 / drawn 2)."""
    _write_fixture_results(tmp_path, with_draw=True)
    games = load_games(tmp_path)
    assert games[2]["tie"] is True and games[2]["won"] is False
    totals = compute_totals(games)
    assert totals == {"legal_games": 3, "points": 172, "won": 1, "lost": 1, "drawn": 1}


def test_missing_ledger_is_a_clear_error(tmp_path):
    with pytest.raises(PdfParserError, match="league ledger"):
        load_games(tmp_path)


def test_unfilled_placeholders_skip_comment_keys():
    data = {
        "_comment": "fill every <FILL:*>",
        "good": "value",
        "bad": "<FILL: me>",
        "nested": {"_note": "<FILL ignored>", "also_bad": "<FILL: too>"},
    }
    found = unfilled_placeholders(data)
    assert sorted(found) == ["bad = <FILL: me>", "nested.also_bad = <FILL: too>"]


def _docx_available():
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed in this venv")
def test_build_submission_fills_template(tmp_path):
    from docx import Document

    from tools.submission_builder import build_submission

    _write_fixture_results(tmp_path)
    template = tmp_path / "template.docx"
    doc = Document()
    for i in range(21):
        if i in (8, 9, 12, 13):
            doc.add_paragraph("First name:   Last name:")
        else:
            doc.add_paragraph(f"Field {i}:")
    doc.add_table(rows=11, cols=9)
    doc.save(str(template))
    data_file = tmp_path / "data.json"
    data_file.write_text(
        json.dumps(
            {
                "group_id": "vibecode",
                "exercise_number": "07",
                "self_score": "90",
                "cop_repo": "https://example/cop",
                "thief_repo": "https://example/thief",
                "agent_email": "agent@example.com",
                "student1": {
                    "id_card": "1",
                    "first_en": "A",
                    "last_en": "K",
                    "first_he": "א",
                    "last_he": "ק",
                },
                "student2": {
                    "id_card": "2",
                    "first_en": "R",
                    "last_en": "M",
                    "first_he": "ר",
                    "last_he": "מ",
                },
                "bonus_eligibility": "No",
                "opponent_agent_emails": {"oppa": "a@x.com", "oppb": "b@x.com"},
            }
        ),
        encoding="utf-8",
    )
    result = build_submission(template, data_file, tmp_path, tmp_path, to_pdf=False)
    assert result["unfilled"] == [] and result["pdf"] is None
    filled = Document(str(result["docx"]))
    assert filled.paragraphs[1].text.endswith("vibecode")
    assert filled.paragraphs[2].text.endswith("90")
    assert "A" in filled.paragraphs[8].text and "K" in filled.paragraphs[8].text
    table = filled.tables[0]
    assert table.rows[1].cells[4].text == "oppa" and table.rows[2].cells[8].text == "b@x.com"
    assert result["docx"].name == "vibecode-ex07.docx"
