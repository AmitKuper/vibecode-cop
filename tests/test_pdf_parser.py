"""tools/pdf_parser — write→read round-trips, search, metadata, CLI, errors."""

from __future__ import annotations

import pytest

from tools.pdf_parser import (
    PdfParserError,
    read_pdf_metadata,
    read_pdf_text,
    search_pdf,
    write_text_pdf,
)
from tools.pdf_parser.cli import main as cli_main
from tools.pdf_parser.writer import CHARS_PER_LINE, LINES_PER_PAGE


def test_write_then_read_round_trip(tmp_path):
    target = tmp_path / "report.pdf"
    written = write_text_pdf(
        target, "Match Report", ["Cop won 90-30.", "Audits: 6/6 Verified OK."]
    )
    assert written == target and target.stat().st_size > 0
    pages = read_pdf_text(target)
    assert len(pages) == 1
    assert "Match Report" in pages[0]
    assert "Cop won 90-30." in pages[0]
    assert "Audits: 6/6 Verified OK." in pages[0]


def test_special_characters_survive_escaping(tmp_path):
    target = tmp_path / "esc.pdf"
    tricky = r"score (90-30) \ backslash and (nested (parens))"
    write_text_pdf(target, "Esc", [tricky])
    assert tricky in read_pdf_text(target)[0]


def test_long_text_flows_onto_multiple_pages(tmp_path):
    target = tmp_path / "long.pdf"
    paragraphs = [f"line {i}" for i in range(LINES_PER_PAGE + 10)]
    write_text_pdf(target, "Long", paragraphs)
    pages = read_pdf_text(target)
    assert len(pages) >= 2
    assert "line 0" in pages[0]
    assert f"line {LINES_PER_PAGE + 9}" in pages[-1]


def test_wrapping_keeps_lines_under_budget(tmp_path):
    target = tmp_path / "wrap.pdf"
    write_text_pdf(target, "Wrap", ["word " * 200])
    text = read_pdf_text(target)[0]
    content_lines = [ln for ln in text.splitlines() if ln.startswith("word")]
    assert content_lines and all(len(ln) <= CHARS_PER_LINE for ln in content_lines)


def test_metadata_and_single_page_selection(tmp_path):
    target = tmp_path / "meta.pdf"
    write_text_pdf(target, "Meta", [f"line {i}" for i in range(LINES_PER_PAGE + 5)])
    meta = read_pdf_metadata(target)
    assert meta["pages"] >= 2
    only_second = read_pdf_text(target, page=2)
    assert len(only_second) == 1 and "Meta" not in only_second[0]
    with pytest.raises(PdfParserError, match="out of range"):
        read_pdf_text(target, page=99)


def test_search_reports_page_and_line(tmp_path):
    target = tmp_path / "search.pdf"
    write_text_pdf(target, "Search", ["alpha", "the NEEDLE sits here", "omega"])
    hits = search_pdf(target, "needle")
    assert len(hits) == 1
    assert hits[0]["page"] == 1 and "NEEDLE" in hits[0]["line"]
    assert search_pdf(target, "NEEDLE", ignore_case=False)[0]["line"].startswith("the NEEDLE")
    assert search_pdf(target, "absent") == []
    with pytest.raises(PdfParserError, match="empty search"):
        search_pdf(target, "")


def test_reader_rejects_missing_and_garbage_files(tmp_path):
    with pytest.raises(PdfParserError, match="not a file"):
        read_pdf_text(tmp_path / "missing.pdf")
    junk = tmp_path / "junk.pdf"
    junk.write_bytes(b"this is not a pdf at all")
    with pytest.raises(PdfParserError, match="unreadable"):
        read_pdf_text(junk)


def test_cli_write_read_meta_search(tmp_path, capsys):
    target = tmp_path / "cli.pdf"
    assert cli_main(["write", str(target), "--title", "CLI", "--text", "hello world"]) == 0
    assert cli_main(["read", str(target)]) == 0
    assert cli_main(["meta", str(target)]) == 0
    assert cli_main(["search", str(target), "hello"]) == 0
    out = capsys.readouterr().out
    assert "hello world" in out and '"pages": 1' in out and "1 hit(s)" in out


def test_cli_write_from_file(tmp_path, capsys):
    notes = tmp_path / "notes.txt"
    notes.write_text("first paragraph\n\nsecond paragraph", encoding="utf-8")
    target = tmp_path / "notes.pdf"
    assert cli_main(["write", str(target), "--title", "Notes", "--from-file", str(notes)]) == 0
    text = read_pdf_text(target)[0]
    assert "first paragraph" in text and "second paragraph" in text


def _docx_available():
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed in this venv")
def test_fill_docx_form_appends_without_moving_fields(tmp_path):
    from docx import Document

    from tools.pdf_parser.docx_form import fill_docx_form

    template = tmp_path / "form.docx"
    doc = Document()
    doc.add_paragraph("Group ID code:")
    doc.add_paragraph("Self score:")
    table = doc.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "opponent"
    table.rows[0].cells[1].text = "score"
    doc.save(str(template))

    out = tmp_path / "filled.docx"
    fill_docx_form(
        template,
        out,
        {0: "vibecode", 1: "90"},
        table_rows=[["imreeyal", "90:30"], ["anrbj666", "35:75"]],
    )
    filled = Document(str(out))
    assert filled.paragraphs[0].text == "Group ID code: vibecode"
    assert filled.paragraphs[1].text == "Self score: 90"
    assert filled.tables[0].rows[1].cells[0].text == "imreeyal"
    assert filled.tables[0].rows[2].cells[1].text == "35:75"


@pytest.mark.skipif(not _docx_available(), reason="python-docx not installed in this venv")
def test_fill_docx_form_rejects_bad_indices(tmp_path):
    from docx import Document

    from tools.pdf_parser.docx_form import fill_docx_form

    template = tmp_path / "form.docx"
    doc = Document()
    doc.add_paragraph("only one line")
    doc.save(str(template))
    with pytest.raises(PdfParserError, match="out of range"):
        fill_docx_form(template, tmp_path / "x.docx", {5: "nope"})
    with pytest.raises(PdfParserError, match="template not found"):
        fill_docx_form(tmp_path / "missing.docx", tmp_path / "x.docx", {0: "v"})
